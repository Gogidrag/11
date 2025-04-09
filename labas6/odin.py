import tkinter as tk
from tkinter import filedialog
from tkinter import messagebox
from vehicles_package.car import Car
from vehicles_package.truck import Truck
from vehicles_package.bus import Bus
from docx import Document

class VehicleApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Расчет поездки")
        self.root.geometry("400x350")

        tk.Label(root, text="Тип транспорта:").grid(row=0, column=0, sticky="w", padx=10, pady=5)
        self.vehicle_type = tk.StringVar(value="Легковой")
        tk.OptionMenu(root, self.vehicle_type, "Легковой", "Грузовой", "Пассажирский").grid(row=0, column=1, padx=10, pady=5)

        tk.Label(root, text="Расстояние (км):").grid(row=1, column=0, sticky="w", padx=10, pady=5)
        self.distance_input = tk.Entry(root)
        self.distance_input.grid(row=1, column=1, padx=10, pady=5)

        tk.Label(root, text="Загрузка:").grid(row=2, column=0, sticky="w", padx=10, pady=5)
        self.load_input = tk.Entry(root)
        self.load_input.grid(row=2, column=1, padx=10, pady=5)

        tk.Button(root, text="Рассчитать", command=self.calculate).grid(row=3, column=0, columnspan=2, pady=10)

        tk.Label(root, text="Результат:").grid(row=4, column=0, sticky="w", padx=10, pady=5)
        self.result_text = tk.Text(root, height=5, width=40)
        self.result_text.grid(row=5, column=0, columnspan=2, padx=10, pady=5)

    def calculate(self):
        try:
            vehicle_type = self.vehicle_type.get()
            distance = float(self.distance_input.get())
            load = float(self.load_input.get())


            if vehicle_type == "Легковой":
                vehicle = Car(fuel_consumption_per_100km=8.8, fuel_price=60.42)
                fuel = vehicle.calculate_fuel_consumption(distance)
                cost = vehicle.calculate_trip_cost(distance)
                time = vehicle.calculate_trip_time(distance, average_speed=60)
            elif vehicle_type == "Грузовой":
                vehicle = Truck(fuel_consumption_per_100km=17, fuel_price=70.43, load_capacity=10)
                fuel = vehicle.calculate_fuel_consumption(distance, load)
                cost = vehicle.calculate_trip_cost(distance, load)
                time = vehicle.calculate_trip_time(distance, average_speed=40)
            elif vehicle_type == "Пассажирский":
                vehicle = Bus(fuel_consumption_per_100km=11, fuel_price=70.47, passenger_capacity=50)
                fuel = vehicle.calculate_fuel_consumption(distance, load)
                cost = vehicle.calculate_trip_cost(distance, load)
                time = vehicle.calculate_trip_time(distance, average_speed=60)
            else:
                raise ValueError("Выберите тип транспорта.")

            result = f"Расход топлива: {fuel:.2f} л\nСтоимость: {cost:.2f} руб.\nВремя: {time:.2f} ч"
            self.result_text.delete(1.0, tk.END)
            self.result_text.insert(tk.END, result)
        except Exception as e:
            messagebox.showerror("Ошибка", str(e))
        self.save_report()
            

    def save_report(self):
        try:
            file_path = filedialog.asksaveasfilename(
                defaultextension=".docx",
                filetypes=[("Word Document", "*.docx")],
                title="Сохранить отчет"
            )
            if not file_path:
                return  
            doc = Document()
            doc.add_heading("Отчет по поездке", level=1)
            for key, value in self.report_data.items():
                doc.add_paragraph(f"{key}: {value}")
            doc.save("C:\Users\Varya\Desktop\labas6\vehicles_package\docx.docx")
            messagebox.showinfo("Готово", f"Отчет сохранен в файл: {file_path}")
        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось сохранить отчет: {e}")

if __name__ == "__main__":
    root = tk.Tk()
    app = VehicleApp(root)
    root.mainloop()